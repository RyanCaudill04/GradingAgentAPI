from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from app.main import app
from app.db.models import Base
from app.api.deps import get_db
from io import BytesIO
import docx
import pytest

print("test_main.py is being executed")

# Setup test database
SQLALCHEMY_DATABASE_URL = "sqlite:///./test.db"
engine = create_engine(
    SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False}
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


@pytest.fixture(name="session")
def session_fixture():
    Base.metadata.create_all(bind=engine)
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()
        Base.metadata.drop_all(bind=engine)


@pytest.fixture(name="client")
def client_fixture(session: TestingSessionLocal):
    def override_get_db():
        yield session

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as client:
        yield client
    app.dependency_overrides.clear()

def test_read_root(client: TestClient):
    response = client.get("/")
    assert response.status_code == 200
    assert response.json() == {"message": "FastAPI is connected!"}

def test_create_assignment(client: TestClient):
    response = client.post("/assignments", json={"assignment_name": "Test Assignment"})
    assert response.status_code == 200
    assert response.json()["name"] == "Test Assignment"
    assert "id" in response.json()

    # Test creating an existing assignment
    response = client.post("/assignments", json={"assignment_name": "Test Assignment"})
    assert response.status_code == 400
    assert response.json() == {"detail": "Assignment already exists"}

def test_upload_criteria_json(client: TestClient):
    assignment_name = "Test Assignment 2"
    client.post("/assignments", json={"assignment_name": assignment_name})
    
    criteria_content = '[{"pattern": "test", "deduction": 10, "message": "test message"}]'
    response = client.post(
        f"/assignments/{assignment_name}/criteria",
        files={
            "criteria_file": (
                "criteria.json",
                criteria_content.encode("utf-8"),
                "application/json"
            )
        }
    )
    assert response.status_code == 200
    assert response.json() == {"message": f"Criteria for {assignment_name} saved."}

def test_upload_criteria_docx(client: TestClient):
    assignment_name = "Test Assignment 3"
    client.post("/assignments", json={"assignment_name": assignment_name})

    document = docx.Document()
    document.add_paragraph('[{"pattern": "test", "deduction": 10, "message": "test message"}]')
    
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    response = client.post(
        f"/assignments/{assignment_name}/criteria",
        files={
            "criteria_file": (
                "criteria.docx",
                file_stream.read(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        }
    )
    assert response.status_code == 200
    assert response.json() == {"message": f"Criteria for {assignment_name} saved."}

def test_upload_criteria_invalid_file_type(client: TestClient):
    assignment_name = "Test Assignment 4"
    client.post("/assignments", json={"assignment_name": assignment_name})

    response = client.post(
        f"/assignments/{assignment_name}/criteria",
        files={
            "criteria_file": (
                "criteria.pdf",
                b"some pdf content",
                "application/pdf"
            )
        }
    )
    assert response.status_code == 400
    assert response.json() == {"detail": "Invalid file type. Only .txt, .docx, and .json files are allowed."}

def test_upload_criteria_docx(client: TestClient):
    assignment_name = "Test Assignment 3"
    client.post("/assignments", json={"assignment_name": assignment_name})

    document = docx.Document()
    document.add_paragraph("This is a test criteria for a docx file.")
    document.add_paragraph("Second paragraph.")
    
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    response = client.post(
        f"/assignments/{assignment_name}/criteria",
        files={
            "criteria_file": (
                "criteria.docx",
                file_stream.read(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        }
    )
    assert response.status_code == 200
    assert response.json() == {"message": f"Criteria for {assignment_name} saved."}

def test_upload_criteria_invalid_file_type(client: TestClient):
    assignment_name = "Test Assignment 4"
    client.post("/assignments", json={"assignment_name": assignment_name})

    response = client.post(
        f"/assignments/{assignment_name}/criteria",
        files={
            "criteria_file": (
                "criteria.pdf",
                b"some pdf content",
                "application/pdf"
            )
        }
    )
    assert response.status_code == 400
    assert response.json() == {"detail": "Invalid file type. Only .txt, .docx, and .json files are allowed."}
